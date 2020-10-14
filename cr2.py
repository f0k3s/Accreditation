import docx
import pandas as pd
import numpy as np

doc = docx.Document()

df = pd.DataFrame()

doc.add_heading('Тест', 2)
doc.add_paragraph('Тесто')
par1 = doc.add_paragraph('Доп.тесто ')
par1.add_run(' Добавить текст')

a=int(input('Если хотите создать таблицу напишитe - 1:' ))
if a > 0:
    row = int(input('Строки: '))
    col = int(input('Столбцы: '))
    table = doc.add_table(rows=row, cols=col)
    table.style = 'Table Grid'
    table.style = 'Table Grid'
else:
    print('Не создаем таблицу')
doc.save('Accr2.docx')



