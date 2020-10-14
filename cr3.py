import docx
import pandas as pd
import numpy as np

MAX_PAGE_SIZE = 80

doc = docx.Document()
d = doc


doc.add_heading('Тест', 2)
doc.add_paragraph('Тесто')
par1 = doc.add_paragraph('Доп.тесто ')
par1.add_run(' Добавить текст')

type_of_table=docx.enum.style.WD_STYLE_TYPE.TABLE

list_table=[['header1','header2'],['cell1','cell2'],['cell3','cell4']]
numcols=max(map(len,list_table))
numrows=len(list_table)

styles=(s for s in d.styles if s.type==type_of_table)

current_page_row = 0

for stylenum,style in enumerate(styles,start=1):
    next_last_row = current_page_row + numrows + 1

    if next_last_row > MAX_PAGE_SIZE:
        d.add_page_break()
        current_page_row = 0

    label=d.add_paragraph('{}) {}'.format(stylenum,style.name))
    current_page_row += 1

    label.paragraph_format.keep_with_next=True
    label.paragraph_format.space_before=docx.shared.Pt(18)
    label.paragraph_format.space_after=docx.shared.Pt(0)

    table=d.add_table(numrows,numcols)
    table.style=style
    for r,row in enumerate(list_table):
        for c,cell in enumerate(row):
            table.row_cells(r)[c].text=cell

    current_page_row += numrows

            
doc.save('Accr.docx')
