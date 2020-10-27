import docx
import pandas as pd
import numpy as np

doc = docx.Document()
df = pd.DataFrame()

###Создать заголовок
def Title():
    TextH=input("Введите заголовок: ")
    StyleH=int(input("Стиль: "))
    doc.add_heading(TextH, StyleH)
Title()

###Новый абзац и добавление текста
def Paragraph():
    ParName=input("Введите текст: ")
    DocName = doc.add_paragraph(ParName)
    addText=input('Добавьте текст: ')
    DocName.add_run(addText)
Paragraph()

###Создание таблицы
def Tablet():
    row = int(input('Строки: '))
    col = int(input('Столбцы: '))
    table = doc.add_table(rows=row, cols=col)
    table.style = 'Table Grid'
Tablet()

doc.save('Testoooo.docx')



